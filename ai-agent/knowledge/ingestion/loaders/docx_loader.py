from typing import List
import docx

from ..document import Document


def load_docx(file_path: str) -> List[Document]:

    doc = docx.Document(file_path)

    docs = []

    for i, para in enumerate(doc.paragraphs):

        text = para.text.strip()

        if not text:
            continue

        docs.append(
            Document(
                page_content=text,
                metadata={
                    "source": file_path,
                    "paragraph": i,
                    "type": "docx"
                }
            )
        )

    return docs